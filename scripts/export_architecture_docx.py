#!/usr/bin/env python3
from __future__ import annotations

import argparse
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_BREAK
except Exception as exc:  # pragma: no cover
    raise SystemExit(
        "python-docx is required. Install with: pip install python-docx"
    ) from exc


def is_heading(line: str) -> int | None:
    """Return heading level (1-6) if the line represents a markdown heading."""
    stripped = line.lstrip()
    if stripped.startswith("#"):
        hashes = len(stripped) - len(stripped.lstrip("#"))
        level = min(max(hashes, 1), 6)
        return level
    return None


def is_underline_heading(prev: str, curr: str) -> int | None:
    """Return heading level for setext headings (=== or ---) based on underline line."""
    # Prev is the text line; curr is the underline line
    line = curr.strip()
    if not prev.strip():
        return None
    if set(line) == {"="}:
        return 1
    if set(line) == {"-"}:
        return 2
    return None


def is_bullet(line: str) -> bool:
    stripped = line.lstrip()
    return stripped.startswith("- ")


def export_markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()

    i = 0
    while i < len(lines):
        line = lines[i]
        # setext heading support
        if i + 1 < len(lines):
            level = is_underline_heading(line, lines[i + 1])
            if level:
                doc.add_heading(line.strip(), level=level)
                i += 2
                continue

        # atx heading support
        level = is_heading(line)
        if level:
            # strip leading '#' and whitespace
            content = line.lstrip("#").strip()
            doc.add_heading(content, level=level)
            i += 1
            continue

        # bullets
        if is_bullet(line):
            # gather consecutive bullets into a single list
            while i < len(lines) and is_bullet(lines[i]):
                item = lines[i].lstrip()
                item = item[2:].rstrip()
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(item)
                i += 1
            continue

        # blank line -> paragraph break
        if not line.strip():
            # create a small break by adding empty paragraph
            doc.add_paragraph("")
            i += 1
            continue

        # normal paragraph
        doc.add_paragraph(line.rstrip())
        i += 1

    # ensure parent directory exists
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Export docs/backend_system_design.md to .docx"
    )
    parser.add_argument(
        "--input",
        type=Path,
        default=Path("docs/backend_system_design.md"),
        help="Input Markdown file path",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("docs/backend_system_design.docx"),
        help="Output DOCX file path",
    )
    args = parser.parse_args()
    export_markdown_to_docx(args.input, args.output)
    print(f"Wrote {args.output}")


if __name__ == "__main__":
    main()


